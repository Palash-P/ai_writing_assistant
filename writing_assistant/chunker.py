# writing_assistant/chunker.py
"""
Smart document chunking — different strategies for different content types.
"""
import logging
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownTextSplitter,
)

logger = logging.getLogger(__name__)


# ── Splitter Configurations ───────────────────────────────────

def get_splitter(content_type='text', chunk_size=1000, chunk_overlap=200):
    """
    Return the right splitter for the content type.

    Why different splitters for different types?
    - Markdown has headers (##, ###) that define natural sections
    - Code has functions/classes that should stay together
    - Plain text has paragraphs separated by blank lines
    - Using the wrong splitter breaks the natural structure
    """
    if content_type == 'markdown':
        # MarkdownTextSplitter splits on headers first (##, ###)
        # keeping sections intact before splitting by size
        return MarkdownTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    elif content_type == 'code':
        # For code, split on class/function definitions first
        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\nclass ",      # class definitions
                "\ndef ",        # function definitions
                "\n\n",          # blank lines
                "\n",            # line breaks
                " ",             # words
                "",              # characters (last resort)
            ]
        )

    else:
        # Default: prose/article text
        # Split on paragraphs first, then sentences, then words
        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n",   # paragraph breaks (best)
                "\n",     # line breaks
                ". ",     # sentence ends
                "? ",     # question ends
                "! ",     # exclamation ends
                "; ",     # semicolons
                ", ",     # commas
                " ",      # words
                "",       # characters (worst case)
            ]
        )


# ── Text Extraction ───────────────────────────────────────────

def extract_from_pdf(file_path):
    """Extract text from PDF with page tracking"""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append({
                'text': text.strip(),
                'page': i + 1,
                'source': 'pdf'
            })

    return pages


def extract_from_docx(file_path):
    """
    Extract text from Word document.
    DOCX files have paragraphs and tables — we extract both.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                full_text.append(row_text)

    # DOCX doesn't have page numbers — treat as one page
    return [{
        'text': '\n\n'.join(full_text),
        'page': 1,
        'source': 'docx'
    }]


def extract_from_txt(file_path):
    """Extract text from plain text file"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    return [{
        'text': content,
        'page': 1,
        'source': 'txt'
    }]


def extract_from_md(file_path):
    """Extract text from Markdown file"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    return [{
        'text': content,
        'page': 1,
        'source': 'markdown'
    }]


def extract_text(file_path):
    """
    Route to the right extractor based on file extension.
    Now supports: PDF, DOCX, TXT, MD, XLSX, XLS, CSV
    """
    import os
    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
    '.pdf':  extract_from_pdf,
    '.docx': extract_from_docx,
    '.doc':  extract_from_docx,
    '.txt':  extract_from_txt,
    '.md':   extract_from_md,
    '.py':   extract_from_python,
    '.js':   extract_from_js,
    '.ts':   extract_from_js,
    '.dart': extract_from_txt,   
}

    if ext in extractors:
        pages = extractors[ext](file_path)
        if not pages or not any(p['text'] for p in pages):
            raise ValueError("No text could be extracted from this file")
        logger.info(f"Extracted {len(pages)} pages from {ext} file")
        return pages, ext

    # Structured data — handled differently (returns chunks directly)
    elif ext in ['.xlsx', '.xls']:
        return None, ext   # signal to process_file to use table processor

    elif ext == '.csv':
        return None, ext

    else:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: .pdf, .docx, .txt, .md, .xlsx, .xls, .csv"
        )


def process_file(file_path, chunk_size=1000, chunk_overlap=200, include_images=True):
    """
    Main entry point — handles all supported file types.
    Returns list of {text, metadata} dicts ready for embedding.
    """
    import os
    from .table_processor import excel_to_chunks, csv_to_chunks

    ext = os.path.splitext(file_path)[1].lower()

    # Handle structured data directly
    if ext in ['.xlsx', '.xls']:
        logger.info("Processing Excel file")
        return excel_to_chunks(file_path)

    if ext == '.csv':
        logger.info("Processing CSV file")
        return csv_to_chunks(file_path)

    # Handle text-based documents
    pages, detected_ext = extract_text(file_path)
    text_chunks = chunk_pages(pages, detected_ext, chunk_size, chunk_overlap)

    # Optionally extract and describe images from PDFs
    if ext == '.pdf' and include_images:
        try:
            from .image_processor import process_pdf_images
            image_chunks = process_pdf_images(file_path)
            if image_chunks:
                logger.info(f"Adding {len(image_chunks)} image chunks")
                text_chunks.extend(image_chunks)
        except Exception as e:
            # Image processing failure shouldn't fail the whole document
            logger.warning(f"Image processing failed (non-fatal): {e}")

    return text_chunks


# ── Smart Chunking ────────────────────────────────────────────

def detect_content_type(ext, text_sample):
    """Detect content type from extension and content sample"""
    ext_map = {
        '.md':   'markdown',
        '.py':   'code',
        '.js':   'code',
        '.ts':   'code',
        '.dart': 'code',
    }

    if ext in ext_map:
        return ext_map[ext]

    # Check if content looks like code
    code_indicators = [
        'def ', 'class ', 'import ', 'from ',  # Python
        'function ', 'const ', 'let ', 'var ',  # JavaScript
        'public ', 'private ', 'void ',         # Java/Dart
    ]
    if any(indicator in text_sample for indicator in code_indicators):
        return 'code'

    return 'text'


def chunk_pages(pages, ext, chunk_size=1000, chunk_overlap=200):
    """
    Chunk extracted pages using the right strategy.
    Returns list of chunk dicts with rich metadata.
    """
    all_chunks = []
    global_chunk_index = 0

    for page_data in pages:
        text = page_data['text']
        page_num = page_data['page']

        # Detect content type for this page
        content_type = detect_content_type(ext, text[:500])
        splitter = get_splitter(content_type, chunk_size, chunk_overlap)

        # Split the page text into chunks
        raw_chunks = splitter.split_text(text)

        for local_index, chunk_text in enumerate(raw_chunks):
            if not chunk_text.strip():
                continue

            all_chunks.append({
                'text': chunk_text.strip(),
                'metadata': {
                    'page': page_num,
                    'chunk_index': global_chunk_index,
                    'local_chunk_index': local_index,
                    'content_type': content_type,
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_text.split()),
                }
            })
            global_chunk_index += 1

    logger.info(f"Created {len(all_chunks)} chunks from {len(pages)} pages")
    return all_chunks


def process_file(file_path, chunk_size=1000, chunk_overlap=200):
    """
    Main entry point — extract and chunk any supported file.
    Returns list of {text, metadata} dicts ready for embedding.
    """
    pages, ext = extract_text(file_path)
    chunks = chunk_pages(pages, ext, chunk_size, chunk_overlap)
    return chunks

def extract_from_python(file_path):
    """Extract Python code with function/class boundaries preserved"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    return [{'text': content, 'page': 1, 'source': 'python'}]


def extract_from_js(file_path):
    """Extract JavaScript/TypeScript code"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    return [{'text': content, 'page': 1, 'source': 'javascript'}]